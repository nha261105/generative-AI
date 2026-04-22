"""
src/application/pipeline_docx.py
─────────────────────────────────
Pipeline xử lý DOCX dùng thuần python-docx (không cần Unstructured).

Luồng: Load DOCX → Extract elements → Enrich metadata → Chunk → Embed → FAISS

Tính năng:
  • Page number chính xác từ XML page break (<w:br w:type="page"/>)
    + lastRenderedPageBreak (Word tự chèn khi render)
  • Parse paragraph (body, heading level 1-9)
  • Parse table → text dạng markdown-style "| col1 | col2 |"
  • Metadata đầy đủ: page, section, heading_level, element_type, source
  • Noise filtering trước khi chunk

Cải thiện retrieve sau này (comment sẵn):
  • Dùng metadata["section"] để filter chunk theo heading
  • Dùng metadata["element_type"] để ưu tiên Table chunk khi query về số liệu
  • Tăng chunk_size lên 1500 nếu document nhiều bảng
  • MMR retriever để tránh retrieve chunks giống nhau
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings

logger = logging.getLogger(__name__)

EMBEDDING_MODEL       = "sentence-transformers/paraphrase-multilingual-mpnet-base-v2"
DEFAULT_CHUNK_SIZE    = 1000
DEFAULT_CHUNK_OVERLAP = 150


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — DOCX READER
# Đọc file DOCX bằng python-docx, trả về list[Document] với metadata đầy đủ
# ══════════════════════════════════════════════════════════════════════════════

def _paragraph_starts_new_page(para) -> bool:
    """
    Trả về True nếu paragraph chứa page break.

    Hai loại page break trong DOCX XML:
      1. <w:br w:type="page"/>       — explicit page break do user chèn
      2. <w:lastRenderedPageBreak/>  — Word tự thêm khi render (soft break)
    """
    p_elem = para._element

    # Explicit page break
    for br in p_elem.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True

    # Soft page break từ Word renderer
    if p_elem.find(".//" + qn("w:lastRenderedPageBreak")) is not None:
        return True

    return False


def _get_heading_level(para) -> Optional[int]:
    """
    Trả về heading level (0–9) nếu paragraph là heading, else None.
      0 = Title style
      1 = Heading 1  (hoặc Subtitle)
      2 = Heading 2
      ...
    """
    style_name = (para.style.name or "").strip()

    if style_name == "Title":
        return 0
    if style_name == "Subtitle":
        return 1

    match = re.match(r"Heading\s+(\d+)", style_name, re.IGNORECASE)
    if match:
        return int(match.group(1))

    return None


def _table_to_text(table) -> str:
    """
    Convert python-docx Table → markdown-style text.

    Ví dụ output:
      | Component | Requirement |
      | RAM | 8 GB minimum |
      | Disk | 10 GB |

    Tại sao markdown-style:
      • LLM hiểu tốt markdown table hơn raw text
      • Giữ được cấu trúc cột → retrieval chính xác hơn khi query về số liệu
    """
    rows_text: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows_text.append("| " + " | ".join(cells) + " |")
    return "\n".join(rows_text)


def _read_docx(docx_path: str, filename: str) -> list[Document]:
    """
    Đọc toàn bộ DOCX bằng python-docx.
    Duyệt body elements theo đúng thứ tự xuất hiện trong file XML
    (paragraph và table đan xen nhau).

    Metadata mỗi Document:
      source        : tên file
      file_type     : "docx"
      page          : số trang (1-indexed, dựa trên XML page break)
      section       : text của heading gần nhất phía trên element này
      element_type  : "heading" | "paragraph" | "table"
      heading_level : chỉ có nếu element_type="heading" (0=Title, 1=H1, ...)
    """
    doc      = DocxDocument(docx_path)
    elements: list[Document] = []
    page     = 1
    section  = "Document Start"

    body = doc.element.body

    for child in body.iterchildren():

        # ── TABLE ─────────────────────────────────────────────────────────────
        if child.tag == qn("w:tbl"):
            tbl_obj = next(
                (t for t in doc.tables if t._element is child), None
            )
            if tbl_obj is None:
                continue

            table_text = _table_to_text(tbl_obj)
            if table_text.strip():
                elements.append(Document(
                    page_content=table_text,
                    metadata={
                        "source":       filename,
                        "file_type":    "docx",
                        "page":         page,
                        "section":      section,
                        "element_type": "table",
                    },
                ))

        # ── PARAGRAPH ─────────────────────────────────────────────────────────
        elif child.tag == qn("w:p"):
            para_obj = next(
                (p for p in doc.paragraphs if p._element is child), None
            )
            if para_obj is None:
                continue

            # Kiểm tra page break TRƯỚC khi lấy text
            if _paragraph_starts_new_page(para_obj):
                page += 1

            text = (para_obj.text or "").strip()
            if not text:
                continue

            heading_level = _get_heading_level(para_obj)

            if heading_level is not None:
                # Heading → cập nhật section tracker
                section = text[:150]
                elements.append(Document(
                    page_content=text,
                    metadata={
                        "source":        filename,
                        "file_type":     "docx",
                        "page":          page,
                        "section":       section,
                        "element_type":  "heading",
                        "heading_level": heading_level,
                    },
                ))
            else:
                elements.append(Document(
                    page_content=text,
                    metadata={
                        "source":       filename,
                        "file_type":    "docx",
                        "page":         page,
                        "section":      section,
                        "element_type": "paragraph",
                    },
                ))

    logger.info(
        f"[DocxReader] '{filename}' → "
        f"{len(elements)} element(s) across {page} page(s)"
    )
    return elements


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — CLEANING
# ══════════════════════════════════════════════════════════════════════════════

_NOISE_PATTERNS = [
    re.compile(r"^\s*\d+\s*$"),            # chỉ là số trang
    re.compile(r"^[-–—=_*•·]{2,}\s*$"),   # divider lines
    re.compile(r"^\s*$"),                   # blank
]


def _is_valid_chunk(text: str) -> bool:
    """
    Lọc chunk quá ngắn hoặc là noise.
    Giữ lại table rows ngắn (| RAM | 8 GB |) vì chứa thông tin cô đọng.
    """
    stripped = text.strip()

    # Table rows: cho qua nếu có cấu trúc | ... |
    if stripped.startswith("|") and stripped.endswith("|"):
        return len(stripped) > 5

    if len(stripped) < 30:
        return False

    for pattern in _NOISE_PATTERNS:
        if pattern.match(stripped):
            return False

    return True


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — MAIN PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

def process_docx_to_vectorstore(
    docx_path: str,
    vector_store_path: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
) -> FAISS:
    """
    Pipeline: Load DOCX → Clean → Chunk → Embed → FAISS.

    Args:
        docx_path:         Đường dẫn đến file .docx
        vector_store_path: Thư mục lưu FAISS index
        chunk_size:        Ký tự tối đa mỗi chunk (default 1000)
        chunk_overlap:     Ký tự trùng lặp giữa chunk liên tiếp (default 150)

    Returns:
        FAISS vector store đã build và lưu xuống disk.

    Raises:
        FileNotFoundError : File không tồn tại.
        ValueError        : File không có text hoặc không tạo được chunk.

    Improvement hooks:
        chunk_size=1500   — nếu DOCX nhiều bảng, cần context rộng hơn
        chunk_size=500    — precision cao cho câu hỏi chi tiết
        MMR retriever     — tránh retrieve chunks giống nhau:
                            vector_db.as_retriever(search_type="mmr")
        Section filter    — chỉ search trong 1 section cụ thể:
                            vector_db.as_retriever(
                                search_kwargs={
                                    "filter": {"section": "Installation Steps"}
                                }
                            )
        Table-only filter — ưu tiên bảng khi query về số liệu:
                            vector_db.as_retriever(
                                search_kwargs={
                                    "filter": {"element_type": "table"}
                                }
                            )
    """
    path     = Path(docx_path)
    filename = path.name

    if not path.exists():
        raise FileNotFoundError(f"Không tìm thấy file: '{docx_path}'")

    logger.info(f"[DocxPipeline] ── Start: '{filename}' ──")

    # ── Step 1: Load & extract ────────────────────────────────────────────────
    logger.info("[DocxPipeline] Step 1/4 — Reading DOCX with python-docx")
    documents = _read_docx(docx_path, filename)

    if not documents:
        raise ValueError(
            f"Không trích xuất được văn bản từ '{filename}'. "
            "File có thể bị hỏng hoặc chỉ chứa hình ảnh."
        )
    logger.info(f"[DocxPipeline] Extracted {len(documents)} element(s)")

    # ── Step 2: Chunk ─────────────────────────────────────────────────────────
    logger.info(
        f"[DocxPipeline] Step 2/4 — Chunking "
        f"(size={chunk_size}, overlap={chunk_overlap})"
    )
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    print("!!!!!!!!!!!!!!!!!!!!!!!!!DOCUMENTS!!!!!!!!!!!!!!!!!!!!!!!")
    for doc in documents:
        print(type(doc))
        print(doc)
    chunks = splitter.split_documents(documents)

    # Filter noise — giống pipeline.py gốc
    chunks = [c for c in chunks if _is_valid_chunk(c.page_content)]

    if not chunks:
        raise ValueError(
            "Không tạo được chunk hợp lệ từ DOCX. "
            "Vui lòng kiểm tra nội dung file."
        )
    logger.info(f"[DocxPipeline] {len(chunks)} valid chunk(s) after filtering")

    # ── Step 3: Embed ─────────────────────────────────────────────────────────
    logger.info("[DocxPipeline] Step 3/4 — Building embeddings")
    embeddings = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

    # ── Step 4: FAISS ─────────────────────────────────────────────────────────
    logger.info("[DocxPipeline] Step 4/4 — Building & saving FAISS index")
    vector_db = FAISS.from_documents(chunks, embeddings)

    Path(vector_store_path).mkdir(parents=True, exist_ok=True)
    vector_db.save_local(vector_store_path)

    logger.info(
        f"[DocxPipeline] ── Done: {len(chunks)} chunks, "
        f"index → '{vector_store_path}' ──"
    )
    return vector_db